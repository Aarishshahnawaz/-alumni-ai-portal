"""
Resume Analyzer Service
Extracts text from PDFs, analyzes skills, and generates ATS scores
"""

import asyncio
import re
from typing import Dict, Any, List, Optional
from loguru import logger
import PyPDF2
import pdfplumber
from docx import Document
from io import BytesIO
import base64

from app.models.schemas import ResumeAnalysisInput, ExtractedSkills, ATSScore
from app.services.model_manager import ModelManager
from app.core.exceptions import FileProcessingException, PredictionException


class ResumeAnalyzerService:
    """Service for analyzing resumes and extracting insights"""
    
    def __init__(self, model_manager: ModelManager):
        self.model_manager = model_manager
        
        # Common resume sections
        self.resume_sections = {
            'contact': ['contact', 'personal', 'info'],
            'summary': ['summary', 'objective', 'profile', 'about'],
            'experience': ['experience', 'work', 'employment', 'career'],
            'education': ['education', 'academic', 'qualification'],
            'skills': ['skills', 'technical', 'competencies', 'expertise'],
            'projects': ['projects', 'portfolio', 'work samples'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'achievements': ['achievements', 'awards', 'honors', 'accomplishments']
        }
    
    async def analyze_resume(self, input_data: ResumeAnalysisInput, file_content: bytes = None) -> Dict[str, Any]:
        """Analyze resume and extract insights"""
        try:
            logger.info(f"Analyzing resume of type: {input_data.file_type}")
            
            # Extract text from file
            extracted_text = await self._extract_text_from_file(
                file_content, input_data.file_type
            )
            
            if not extracted_text or len(extracted_text.strip()) < 50:
                raise FileProcessingException("Could not extract sufficient text from file")
            
            # Extract skills
            skills = await self._extract_skills(extracted_text)
            
            # Calculate ATS score
            ats_score = await self._calculate_ats_score(
                extracted_text, skills, input_data.target_role
            )
            
            # Calculate job match score if target role provided
            job_match_score = None
            if input_data.target_role:
                job_match_score = await self._calculate_job_match_score(
                    skills, input_data.target_role
                )
            
            # Generate improvement suggestions
            improvements = await self._generate_improvements(
                extracted_text, skills, ats_score, input_data.target_role
            )
            
            return {
                'extracted_text': extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
                'skills': skills,
                'ats_score': ats_score,
                'job_match_score': job_match_score,
                'improvements': improvements,
                'analysis': {
                    'text_length': len(extracted_text),
                    'word_count': len(extracted_text.split()),
                    'sections_detected': await self._detect_sections(extracted_text),
                    'readability_score': await self._calculate_readability_score(extracted_text)
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing resume: {e}")
            if isinstance(e, (FileProcessingException, PredictionException)):
                raise
            raise PredictionException(str(e))
    
    async def _extract_text_from_file(self, file_content: bytes, file_type: str) -> str:
        """Extract text from different file types"""
        try:
            if file_type.lower() == 'pdf':
                return await self._extract_text_from_pdf(file_content)
            elif file_type.lower() in ['doc', 'docx']:
                return await self._extract_text_from_docx(file_content)
            elif file_type.lower() == 'txt':
                return file_content.decode('utf-8')
            else:
                raise FileProcessingException(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_type}: {e}")
            raise FileProcessingException(f"Failed to extract text: {str(e)}")
    
    async def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text
                
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        try:
            # Method 2: Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    
        except Exception as e:
            logger.error(f"PyPDF2 also failed: {e}")
            raise FileProcessingException("Could not extract text from PDF")
        
        if not text.strip():
            raise FileProcessingException("PDF appears to be empty or contains only images")
        
        return text
    
    async def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise FileProcessingException("Could not extract text from DOCX file")
    
    async def _extract_skills(self, text: str) -> ExtractedSkills:
        """Extract different types of skills from resume text"""
        try:
            text_lower = text.lower()
            
            # Get skill databases
            technical_skills_db = self.model_manager.get_model('technical_skills')
            soft_skills_db = self.model_manager.get_model('soft_skills')
            
            # Extract technical skills
            technical_skills = []
            for skill in technical_skills_db:
                if skill in text_lower:
                    technical_skills.append(skill.title())
            
            # Extract soft skills
            soft_skills = []
            for skill in soft_skills_db:
                if skill in text_lower:
                    soft_skills.append(skill.title())
            
            # Extract certifications using patterns
            certifications = await self._extract_certifications(text)
            
            # Extract tools and technologies
            tools = await self._extract_tools(text)
            
            return ExtractedSkills(
                technical_skills=list(set(technical_skills)),
                soft_skills=list(set(soft_skills)),
                certifications=certifications,
                tools=tools
            )
            
        except Exception as e:
            logger.error(f"Error extracting skills: {e}")
            return ExtractedSkills()
    
    async def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume text"""
        certifications = []
        
        # Common certification patterns
        cert_patterns = [
            r'AWS\s+Certified\s+[\w\s]+',
            r'Microsoft\s+Certified\s+[\w\s]+',
            r'Google\s+Cloud\s+[\w\s]+',
            r'Cisco\s+Certified\s+[\w\s]+',
            r'CompTIA\s+[\w\s]+',
            r'PMP\s+Certified',
            r'Scrum\s+Master\s+Certified',
            r'Six\s+Sigma\s+[\w\s]+',
            r'CISSP',
            r'CISA',
            r'CISM'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        return list(set(certifications))
    
    async def _extract_tools(self, text: str) -> List[str]:
        """Extract tools and technologies from resume text"""
        tools_db = [
            'Git', 'Docker', 'Kubernetes', 'Jenkins', 'Jira', 'Confluence',
            'Slack', 'Trello', 'Asana', 'Figma', 'Adobe Creative Suite',
            'Photoshop', 'Illustrator', 'InDesign', 'Sketch', 'InVision',
            'Tableau', 'Power BI', 'Excel', 'Google Analytics', 'Salesforce'
        ]
        
        text_lower = text.lower()
        found_tools = []
        
        for tool in tools_db:
            if tool.lower() in text_lower:
                found_tools.append(tool)
        
        return list(set(found_tools))
    
    async def _calculate_ats_score(self, text: str, skills: ExtractedSkills, target_role: Optional[str]) -> ATSScore:
        """Calculate ATS (Applicant Tracking System) score"""
        try:
            # Initialize scores
            keyword_score = 0
            format_score = 0
            content_score = 0
            suggestions = []
            
            # Keyword density analysis
            total_skills = len(skills.technical_skills) + len(skills.soft_skills)
            word_count = len(text.split())
            
            if word_count > 0:
                keyword_density = total_skills / word_count
            else:
                keyword_density = 0
            
            # Keyword scoring
            if total_skills >= 15:
                keyword_score = 90
            elif total_skills >= 10:
                keyword_score = 75
            elif total_skills >= 5:
                keyword_score = 60
            else:
                keyword_score = 40
                suggestions.append("Add more relevant technical and soft skills")
            
            # Format scoring
            format_score = await self._calculate_format_score(text, suggestions)
            
            # Content scoring
            content_score = await self._calculate_content_score(text, skills, suggestions)
            
            # Overall score (weighted average)
            overall_score = int(
                keyword_score * 0.4 +
                format_score * 0.3 +
                content_score * 0.3
            )
            
            return ATSScore(
                overall_score=overall_score,
                keyword_density=round(keyword_density, 4),
                format_score=format_score,
                content_score=content_score,
                suggestions=suggestions
            )
            
        except Exception as e:
            logger.error(f"Error calculating ATS score: {e}")
            return ATSScore(
                overall_score=50,
                keyword_density=0.0,
                format_score=50,
                content_score=50,
                suggestions=["Error calculating ATS score"]
            )
    
    async def _calculate_format_score(self, text: str, suggestions: List[str]) -> int:
        """Calculate format score for ATS compatibility"""
        score = 100
        
        # Check for common formatting issues
        if len(text) < 300:
            score -= 20
            suggestions.append("Resume appears too short - add more detail")
        
        if len(text) > 4000:
            score -= 10
            suggestions.append("Resume may be too long - consider condensing")
        
        # Check for section headers
        sections_found = 0
        for section_type, keywords in self.resume_sections.items():
            for keyword in keywords:
                if keyword in text.lower():
                    sections_found += 1
                    break
        
        if sections_found < 4:
            score -= 15
            suggestions.append("Add clear section headers (Experience, Education, Skills, etc.)")
        
        # Check for contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'[\+]?[1-9]?[0-9]{7,15}'
        
        if not re.search(email_pattern, text):
            score -= 10
            suggestions.append("Include a professional email address")
        
        if not re.search(phone_pattern, text):
            score -= 5
            suggestions.append("Include a phone number")
        
        return max(0, score)
    
    async def _calculate_content_score(self, text: str, skills: ExtractedSkills, suggestions: List[str]) -> int:
        """Calculate content quality score"""
        score = 100
        
        # Check for quantifiable achievements
        number_pattern = r'\b\d+%|\b\d+\s*(million|thousand|k|m)\b|\$\d+|\b\d+\s*years?\b'
        numbers_found = len(re.findall(number_pattern, text, re.IGNORECASE))
        
        if numbers_found < 3:
            score -= 15
            suggestions.append("Add quantifiable achievements (percentages, dollar amounts, timeframes)")
        
        # Check for action verbs
        action_verbs = [
            'achieved', 'developed', 'implemented', 'managed', 'led', 'created',
            'improved', 'increased', 'reduced', 'optimized', 'designed', 'built'
        ]
        
        action_verbs_found = sum(1 for verb in action_verbs if verb in text.lower())
        
        if action_verbs_found < 5:
            score -= 10
            suggestions.append("Use more action verbs to describe your accomplishments")
        
        # Check for education section
        education_keywords = ['university', 'college', 'degree', 'bachelor', 'master', 'phd']
        education_found = any(keyword in text.lower() for keyword in education_keywords)
        
        if not education_found:
            score -= 10
            suggestions.append("Include education information")
        
        # Check skill diversity
        if len(skills.technical_skills) < 5:
            score -= 10
            suggestions.append("Add more technical skills relevant to your field")
        
        if len(skills.soft_skills) < 3:
            score -= 5
            suggestions.append("Include relevant soft skills")
        
        return max(0, score)
    
    async def _calculate_job_match_score(self, skills: ExtractedSkills, target_role: str) -> float:
        """Calculate how well the resume matches a target job role"""
        try:
            # Get career data for role matching
            career_data = self.model_manager.get_model('career_data')
            
            # Normalize target role
            target_role_normalized = target_role.lower().replace(' ', '_')
            
            if target_role_normalized not in career_data:
                # Try to find similar role
                for role_name in career_data.keys():
                    if any(word in role_name for word in target_role.lower().split()):
                        target_role_normalized = role_name
                        break
            
            if target_role_normalized not in career_data:
                return 0.5  # Default score if role not found
            
            required_skills = career_data[target_role_normalized]['skills']
            resume_skills = [skill.lower() for skill in skills.technical_skills + skills.soft_skills]
            
            # Calculate skill overlap
            skill_overlap = len(set(required_skills).intersection(set(resume_skills)))
            total_required = len(required_skills)
            
            match_score = skill_overlap / total_required if total_required > 0 else 0
            
            return round(match_score, 3)
            
        except Exception as e:
            logger.error(f"Error calculating job match score: {e}")
            return 0.0
    
    async def _generate_improvements(self, text: str, skills: ExtractedSkills, 
                                   ats_score: ATSScore, target_role: Optional[str]) -> List[str]:
        """Generate improvement suggestions for the resume"""
        improvements = []
        
        # Add ATS suggestions
        improvements.extend(ats_score.suggestions)
        
        # Role-specific improvements
        if target_role:
            career_data = self.model_manager.get_model('career_data')
            target_role_normalized = target_role.lower().replace(' ', '_')
            
            if target_role_normalized in career_data:
                required_skills = career_data[target_role_normalized]['skills']
                resume_skills = [skill.lower() for skill in skills.technical_skills]
                missing_skills = set(required_skills) - set(resume_skills)
                
                if missing_skills:
                    improvements.append(f"Consider adding these skills for {target_role}: {', '.join(list(missing_skills)[:5])}")
        
        # General improvements
        if len(skills.certifications) == 0:
            improvements.append("Consider adding relevant certifications to strengthen your profile")
        
        if len(text.split()) < 200:
            improvements.append("Expand your resume with more detailed descriptions of your experience")
        
        # Remove duplicates and limit to top 10
        improvements = list(dict.fromkeys(improvements))[:10]
        
        return improvements
    
    async def _detect_sections(self, text: str) -> List[str]:
        """Detect which resume sections are present"""
        detected_sections = []
        text_lower = text.lower()
        
        for section_type, keywords in self.resume_sections.items():
            for keyword in keywords:
                if keyword in text_lower:
                    detected_sections.append(section_type)
                    break
        
        return detected_sections
    
    async def _calculate_readability_score(self, text: str) -> float:
        """Calculate a simple readability score"""
        try:
            sentences = len(re.split(r'[.!?]+', text))
            words = len(text.split())
            
            if sentences == 0:
                return 0.0
            
            avg_sentence_length = words / sentences
            
            # Simple readability score (lower is better for resumes)
            # Ideal sentence length for resumes is 15-20 words
            if 15 <= avg_sentence_length <= 20:
                return 1.0
            elif 10 <= avg_sentence_length <= 25:
                return 0.8
            elif 8 <= avg_sentence_length <= 30:
                return 0.6
            else:
                return 0.4
                
        except Exception:
            return 0.5